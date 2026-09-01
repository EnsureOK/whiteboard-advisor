"""知识库管线:上传 -> 解析 -> 切分 -> 向量化入库 -> 混合检索。

- 解析: pypdf / python-docx / 纯文本 & markdown & html
- 切分: 段落优先 + 最大长度控制
- 向量: 千帆 embeddings;无 key 时用确定性 mock 向量,流程同样可跑通
- 关键词召回: Python 侧子串计数(库存量小,无需 FTS 虚表),与向量分数合并

注:检索走"整表读入 + Python 过滤/打分",本地单机数据量(万级 chunk)内足够快,
也避免复杂 JOIN 查询。
"""

from __future__ import annotations

import html as html_mod
import json
import logging
import os
import re
from typing import Optional

import numpy as np
from sqlalchemy.orm import Session as OrmSession

from app.db_models import KbChunk, KbDocument, utcnow_iso
from app.services import embedding as emb

logger = logging.getLogger("whiteboard-advisor.kb")

from app.paths import data_path

UPLOAD_DIR = data_path("uploads")
os.makedirs(UPLOAD_DIR, exist_ok=True)

CHUNK_SIZE = 500       # 单 chunk 目标长度(字符)
CHUNK_OVERLAP = 80     # 相邻 chunk 重叠
MAX_CHUNKS_PER_DOC = 800

SUPPORTED_EXT = {".pdf": "pdf", ".docx": "docx", ".txt": "txt", ".md": "md", ".html": "html", ".htm": "html"}

# 所有外键 id 均为 32 位 hex(uuid4().hex),入口处强校验
_ID_RE = re.compile(r"^[0-9a-f]{32}$")


def check_id(value: str, what: str = "id") -> str:
    if not _ID_RE.fullmatch(value or ""):
        raise ValueError(f"invalid {what}")
    return value


def client_scope(client_id: str) -> str:
    return f"client:{client_id}"


def safe_upload_path(doc_id: str, filename: str) -> str:
    """把上传文件落盘路径限制在 UPLOAD_DIR 内。

    doc_id 已过 check_id;文件名取 basename 并过滤可疑字符,
    最终 normalize 后必须仍位于 UPLOAD_DIR 之内。
    """
    check_id(doc_id, "document id")
    name = os.path.basename(filename or "untitled").replace("..", "_")
    if not name or name in {".", "/"}:
        name = "untitled"
    path = os.path.abspath(os.path.join(UPLOAD_DIR, doc_id + "_" + name))
    if not path.startswith(os.path.abspath(UPLOAD_DIR) + os.sep):
        raise ValueError("invalid upload path")
    return path


def save_upload_file(doc: KbDocument, raw: bytes) -> str:
    """把上传内容写入 UPLOAD_DIR。路径由 safe_upload_path 规范化并校验目录包含。"""
    from pathlib import Path

    path = safe_upload_path(doc.id, doc.filename or "untitled")
    if not os.path.abspath(path).startswith(os.path.abspath(UPLOAD_DIR) + os.sep):
        raise ValueError("invalid upload path")
    Path(path).write_bytes(raw)
    return path


# ---------- 文本抽取 ----------

def extract_text(doc_type: str, raw: bytes) -> str:
    if doc_type == "pdf":
        from io import BytesIO

        from pypdf import PdfReader

        reader = PdfReader(BytesIO(raw))
        return "\n".join(page.extract_text() or "" for page in reader.pages)
    if doc_type == "docx":
        from io import BytesIO

        import docx  # python-docx

        document = docx.Document(BytesIO(raw))
        return "\n".join(p.text for p in document.paragraphs if p.text.strip())
    if doc_type == "html":
        text = raw.decode("utf-8", errors="replace")
        text = re.sub(r"<(script|style)[^>]*>.*?</\1>", " ", text, flags=re.S | re.I)
        text = re.sub(r"<[^>]+>", "\n", text)
        return html_mod.unescape(text)
    # txt / md / text
    return raw.decode("utf-8", errors="replace")


# ---------- 切分 ----------

def chunk_text(text: str) -> list[str]:
    text = text.strip()
    if not text:
        return []
    paras = [p.strip() for p in re.split(r"\n\s*\n", text) if p.strip()]
    pieces: list[str] = []
    for para in paras:
        if len(para) <= CHUNK_SIZE:
            pieces.append(para)
            continue
        # 长段落:先按句子边界,单句超长再硬切
        sentences: list[str] = []
        for s in re.split(r"(?<=[。!?;;\.\n])", para):
            if len(s) > CHUNK_SIZE:
                sentences.extend(s[i : i + CHUNK_SIZE] for i in range(0, len(s), CHUNK_SIZE))
            elif s:
                sentences.append(s)
        buf = ""
        for s in sentences:
            if len(buf) + len(s) > CHUNK_SIZE and buf:
                pieces.append(buf)
                buf = buf[-CHUNK_OVERLAP:] if CHUNK_OVERLAP < len(buf) else buf
            buf += s
        if buf.strip():
            pieces.append(buf)

    # 相邻短片段合并,减少碎片
    merged: list[str] = []
    for p in pieces:
        if merged and len(merged[-1]) + len(p) < CHUNK_SIZE // 2:
            merged[-1] = merged[-1] + "\n" + p
        else:
            merged.append(p)
    return merged[:MAX_CHUNKS_PER_DOC]


# ---------- 入库 ----------

def create_document(
    db: OrmSession,
    *,
    title: str,
    doc_type: str,
    raw: bytes = b"",
    filename: str = "",
    tags: Optional[list[str]] = None,
    scope: str = "global",
    source_url: Optional[str] = None,
) -> KbDocument:
    doc = KbDocument(
        title=title[:200],
        filename=filename[:300],
        doc_type=doc_type,
        size_bytes=len(raw),
        tags_json=json.dumps(tags or [], ensure_ascii=False),
        scope=scope,
        status="uploaded",
        source_url=source_url,
    )
    db.add(doc)
    db.commit()
    db.refresh(doc)
    return doc


def _write_chunks(db: OrmSession, doc: KbDocument, chunks: list[str], vecs: list[list[float]]) -> None:
    db.query(KbChunk).filter(KbChunk.doc_id == doc.id).delete()
    for i, (chunk, vec) in enumerate(zip(chunks, vecs)):
        arr = np.asarray(vec, dtype=np.float32)
        db.add(
            KbChunk(
                doc_id=doc.id,
                seq=i,
                text=chunk,
                embedding=arr.tobytes(),
                dim=int(arr.shape[0]),
            )
        )
    doc.chunk_count = len(chunks)
    doc.status = "indexed"
    doc.error = ""
    doc.updated_at = utcnow_iso()


async def _fetch_url(url: str) -> tuple[str, bytes]:
    import httpx

    async with httpx.AsyncClient(timeout=30, follow_redirects=True) as client:
        resp = await client.get(url, headers={"User-Agent": "Mozilla/5.0 (compatible; WorkbenchKB/0.1)"})
        resp.raise_for_status()
        return resp.headers.get("content-type") or "", resp.content


async def index_document(db: OrmSession, doc_id: str) -> KbDocument:
    """解析 + 切分 + 向量化入库。状态机: parsing -> indexed / failed。"""
    check_id(doc_id, "document id")
    doc = db.get(KbDocument, doc_id)
    if not doc:
        raise ValueError("document not found")
    try:
        doc.status = "parsing"
        doc.error = ""
        db.commit()

        if doc.source_url:
            content_type, raw = await _fetch_url(doc.source_url)
            doc_type = "html" if "html" in content_type else "text"
            text = extract_text(doc_type, raw)
        else:
            path = safe_upload_path(doc.id, doc.filename or "f")
            with open(path, "rb") as f:
                raw = f.read()
            text = extract_text(doc.doc_type, raw)

        chunks = chunk_text(text)
        if not chunks:
            raise ValueError("未解析出有效文本内容")

        vecs = await emb.embed_texts(chunks)
        _write_chunks(db, doc, chunks, vecs)
        db.commit()
        db.refresh(doc)
        return doc
    except Exception as e:  # noqa: BLE001 状态机要求兜底
        logger.warning("kb index failed for %s: %s", doc_id, e)
        doc = db.get(KbDocument, doc_id)
        if doc:
            doc.status = "failed"
            doc.error = str(e)[:500]
            db.commit()
        raise


async def index_inline_text(db: OrmSession, doc: KbDocument, text: str) -> KbDocument:
    """手动粘贴的文本直接入库(不落文件)。"""
    try:
        doc.status = "parsing"
        db.commit()
        chunks = chunk_text(text)
        if not chunks:
            raise ValueError("内容为空")
        vecs = await emb.embed_texts(chunks)
        _write_chunks(db, doc, chunks, vecs)
        db.commit()
        db.refresh(doc)
        return doc
    except Exception as e:  # noqa: BLE001
        doc.status = "failed"
        doc.error = str(e)[:500]
        db.commit()
        raise


# ---------- 检索 ----------

def _keyword_hits(query: str, text: str) -> float:
    """中文按 2-gram 统计命中密度,返回 0~1 的关键词分。"""
    q = re.sub(r"[^\w\u4e00-\u9fff]+", " ", query).strip()
    if not q:
        return 0.0
    grams: list[str] = []
    for seg in q.split():
        if re.fullmatch(r"[\u4e00-\u9fff]+", seg):
            if len(seg) <= 2:
                grams.append(seg)
            else:
                grams.extend(seg[i : i + 2] for i in range(0, len(seg) - 1))
        else:
            grams.append(seg.lower())
    if not grams:
        return 0.0
    hits = sum(1 for g in set(grams) if g in text.lower())
    return min(1.0, hits / max(1, len(set(grams))))


def _visible_rows(db: OrmSession, client_id: Optional[str]) -> list[tuple[KbChunk, KbDocument]]:
    """当前可见的 (chunk, doc) 行:全局库 + 指定客户私有库,已入库状态。

    整表读入后 Python 过滤/关联;向量打分本来就要全量加载 chunk。
    """
    doc_by_id = {d.id: d for d in db.query(KbDocument).all()}
    allowed_scope = client_scope(client_id) if client_id else None
    out: list[tuple[KbChunk, KbDocument]] = []
    for chunk in db.query(KbChunk).all():
        doc = doc_by_id.get(chunk.doc_id)
        if doc is None or doc.status != "indexed":
            continue
        if doc.scope == "global" or (allowed_scope and doc.scope == allowed_scope):
            out.append((chunk, doc))
    return out


async def search_async(
    db: OrmSession,
    query: str,
    *,
    client_id: Optional[str] = None,
    top_k: int = 6,
) -> list[dict]:
    """混合检索: 向量余弦(0.7) + 关键词命中(0.3),返回带来源的 chunk 列表。

    client_id 为 None 时只搜全局库;给出时同时搜该客户私有库。
    """
    if not query.strip():
        return []
    if client_id:
        check_id(client_id, "client id")

    rows = _visible_rows(db, client_id)
    if not rows:
        return []

    # 向量召回(真实或 mock 向量均参与;维度不一致的条目自动跳过)
    [qvec] = await emb.embed_texts([query])
    cand = [(chunk.id, chunk.embedding) for chunk, _ in rows if chunk.embedding]
    vec_scores = dict(emb.cosine_topk(qvec, cand, top_k=top_k * 2))

    combined: dict[str, float] = {}
    for chunk, _ in rows:
        kw = _keyword_hits(query, chunk.text)
        s = 0.7 * vec_scores.get(chunk.id, 0.0) + 0.3 * kw
        if s > 0:
            combined[chunk.id] = s
    ranked = sorted(combined.items(), key=lambda x: -x[1])[:top_k]

    info = {chunk.id: (chunk, doc) for chunk, doc in rows}
    out = []
    for cid, score in ranked:
        chunk, doc = info[cid]
        out.append(
            {
                "chunkId": chunk.id,
                "docId": doc.id,
                "docTitle": doc.title,
                "docType": doc.doc_type,
                "scope": doc.scope,
                "text": chunk.text,
                "score": round(score, 4),
            }
        )
    return out


def delete_document(db: OrmSession, doc_id: str) -> bool:
    check_id(doc_id, "document id")
    doc = db.get(KbDocument, doc_id)
    if not doc:
        return False
    if doc.filename:
        try:
            path = safe_upload_path(doc.id, doc.filename)
        except ValueError:
            path = ""
        if path and os.path.exists(path):
            try:
                os.remove(path)
            except OSError:
                pass
    db.delete(doc)
    db.commit()
    return True
